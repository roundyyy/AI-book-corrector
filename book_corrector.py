#!/usr/bin/env python3
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import os
import requests
import random
import string
import tiktoken
import json

# For reading/writing .docx
try:
    import docx
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    docx = None

LANGUAGES = [
    "English (British)",  # default
    "Albanian",
    "Amharic",
    "Arabic",
    "Armenian",
    "Bengali",
    "Bosnian",
    "Bulgarian",
    "Burmese",
    "Catalan",
    "Chinese",
    "Croatian",
    "Czech",
    "Danish",
    "Dutch",
    "Estonian",
    "Finnish",
    "French",
    "Georgian",
    "German",
    "Greek",
    "Gujarati",
    "Hindi",
    "Hungarian",
    "Icelandic",
    "Indonesian",
    "Italian",
    "Japanese",
    "Kannada",
    "Kazakh",
    "Korean",
    "Latvian",
    "Lithuanian",
    "Macedonian",
    "Malay",
    "Malayalam",
    "Marathi",
    "Mongolian",
    "Norwegian",
    "Persian",
    "Polish",
    "Portuguese",
    "Punjabi",
    "Romanian",
    "Russian",
    "Serbian",
    "Slovak",
    "Slovenian",
    "Somali",
    "Spanish",
    "Swahili",
    "Swedish",
    "Tagalog",
    "Tamil",
    "Telugu",
    "Thai",
    "Turkish",
    "Ukrainian",
    "Urdu",
    "Vietnamese",
]

CHAT_COMPLETIONS_URL = "https://api.openai.com/v1/chat/completions"


class AICorrectorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("AI Book Corrector (DOCX only)")

        # Variables
        self.input_file_path = tk.StringVar()
        self.output_file_path = tk.StringVar()
        self.selected_language = tk.StringVar(value="English (British)")
        self.fix_grammar = tk.BooleanVar(value=True)
        self.fix_typos = tk.BooleanVar(value=True)
        self.fix_formatting = tk.BooleanVar(value=True)
        self.chunk_size = tk.IntVar(value=5000)  # default 5k
        self.extra_instructions = tk.StringVar(value="")

        self.api_key_var = tk.StringVar(value="REPLACE_WITH_YOUR_ACTUAL_KEY")

        self._processing_thread = None
        self._stop_requested = False

        self.build_ui()

    def build_ui(self):
        # == API Key ==
        frame_api = ttk.LabelFrame(self.root, text="API Key")
        frame_api.pack(fill="x", padx=10, pady=5)
        ttk.Label(frame_api, text="OpenAI Key:").pack(side="left", padx=5)
        entry_api = ttk.Entry(
            frame_api, textvariable=self.api_key_var, width=70, show="*")
        entry_api.pack(side="left", padx=5)

        # Right-click paste menu
        self.create_context_menu(entry_api)

        # == Input file (DOCX only) ==
        frame_file = ttk.LabelFrame(
            self.root, text="Select Input File (.docx only)")
        frame_file.pack(fill="x", padx=10, pady=5)
        btn_browse = ttk.Button(
            frame_file, text="Browse...", command=self.browse_file)
        btn_browse.pack(side="left", padx=5, pady=5)
        lbl_input = ttk.Label(
            frame_file, textvariable=self.input_file_path, width=60)
        lbl_input.pack(side="left", padx=5)

        # == Output file (DOCX only) ==
        frame_output = ttk.LabelFrame(
            self.root, text="Select Output File (.docx only)")
        frame_output.pack(fill="x", padx=10, pady=5)
        btn_save = ttk.Button(
            frame_output, text="Save As...", command=self.save_file)
        btn_save.pack(side="left", padx=5, pady=5)
        lbl_output = ttk.Label(
            frame_output, textvariable=self.output_file_path, width=60)
        lbl_output.pack(side="left", padx=5)

        # == Language ==
        frame_language = ttk.LabelFrame(self.root, text="Language Settings")
        frame_language.pack(fill="x", padx=10, pady=5)
        ttk.Label(frame_language, text="Language:").pack(side="left", padx=5)
        combo_lang = ttk.Combobox(
            frame_language,
            textvariable=self.selected_language,
            values=LANGUAGES,
            state="readonly",
            width=30
        )
        combo_lang.pack(side="left", padx=5)

        # == Correction Options ==
        frame_options = ttk.LabelFrame(self.root, text="Correction Options")
        frame_options.pack(fill="x", padx=10, pady=5)
        ttk.Checkbutton(frame_options, text="Fix Grammar",
                        variable=self.fix_grammar).pack(side="left", padx=5)
        ttk.Checkbutton(frame_options, text="Fix Typos",
                        variable=self.fix_typos).pack(side="left", padx=5)
        ttk.Checkbutton(frame_options, text="Fix Formatting",
                        variable=self.fix_formatting).pack(side="left", padx=5)

        # == Chunk Size ==
        frame_chunk = ttk.LabelFrame(self.root, text="Chunk Size (tokens)")
        frame_chunk.pack(fill="x", padx=10, pady=5)
        ttk.Label(frame_chunk, text="Max tokens:").pack(side="left", padx=5)
        spin_chunk = ttk.Spinbox(
            frame_chunk, from_=500, to=32768, textvariable=self.chunk_size, width=8)
        spin_chunk.pack(side="left", padx=5)

        # == Extra Instructions ==
        frame_instructions = ttk.LabelFrame(
            self.root, text="Extra Instructions")
        frame_instructions.pack(fill="both", expand=True, padx=10, pady=5)
        text_instructions = tk.Text(frame_instructions, wrap="word", height=5)
        text_instructions.pack(fill="both", expand=True, padx=5, pady=5)

        def on_text_change(event=None):
            self.extra_instructions.set(text_instructions.get("1.0", "end-1c"))
        text_instructions.bind("<<Modified>>", lambda e: on_text_change())

        # == Action Buttons ==
        frame_actions = ttk.Frame(self.root)
        frame_actions.pack(fill="x", padx=10, pady=10)
        btn_run = ttk.Button(
            frame_actions, text="Run Correction", command=self.start_thread)
        btn_run.pack(side="left", padx=5)
        btn_quit = ttk.Button(frame_actions, text="Quit",
                              command=self.root.quit)
        btn_quit.pack(side="right", padx=5)

    def create_context_menu(self, widget):
        """Right-click context menu with paste."""
        menu = tk.Menu(widget, tearoff=0)

        def do_paste():
            try:
                txt = widget.selection_get(selection='CLIPBOARD')
                widget.insert('insert', txt)
            except tk.TclError:
                pass
        menu.add_command(label="Paste", command=do_paste)

        def show_menu(event):
            menu.tk_popup(event.x_root, event.y_root)
        widget.bind("<Button-3>", show_menu)

    ##############################
    # File selection
    ##############################
    def browse_file(self):
        file_path = filedialog.askopenfilename(
            title="Select input file (DOCX only)",
            filetypes=[("DOCX files", "*.docx")],
        )
        if file_path:
            self.input_file_path.set(file_path)

    def save_file(self):
        file_path = filedialog.asksaveasfilename(
            title="Select output file (DOCX only)",
            defaultextension=".docx",
            filetypes=[("DOCX files", "*.docx")],
        )
        if file_path:
            self.output_file_path.set(file_path)

    ##############################
    # Thread management
    ##############################
    def start_thread(self):
        if self._processing_thread and self._processing_thread.is_alive():
            messagebox.showinfo("Info", "Already running. Please wait...")
            return
        self._stop_requested = False
        self._processing_thread = threading.Thread(
            target=self.run_correction, daemon=True)
        self._processing_thread.start()

    def run_correction(self):
        in_path = self.input_file_path.get().strip()
        out_path = self.output_file_path.get().strip()

        # Validate .docx
        if not in_path.lower().endswith(".docx") or not os.path.isfile(in_path):
            messagebox.showerror(
                "Error", "Please select a valid .docx input file.")
            return
        if not out_path.lower().endswith(".docx"):
            messagebox.showerror(
                "Error", "Please select a valid .docx output file.")
            return

        api_key = self.api_key_var.get().strip()
        if not api_key.startswith("sk-"):
            messagebox.showerror(
                "Error", "Please provide a valid GPT-4o-mini API key.")
            return

        print("[DEBUG] Starting run_correction...")
        paragraphs = []

        # Read docx
        try:
            paragraphs = self.read_docx(in_path)
        except Exception as e:
            messagebox.showerror(
                "Error", f"Failed to read input docx.\n{str(e)}")
            return

        if not paragraphs:
            messagebox.showinfo("Info", "No text found in the .docx file.")
            return

        all_chunks = self.create_chunks(paragraphs, self.chunk_size.get())
        print(f"[DEBUG] Created {len(all_chunks)} chunks for correction.")

        # Always do JSON approach
        merged_paragraphs = []
        for i, chunk in enumerate(all_chunks, start=1):
            if self._stop_requested:
                break
            print(f"[DEBUG] Processing chunk {i}/{len(all_chunks)}")
            partial_json = self.correct_chunk_json_mode(chunk, api_key)
            if partial_json and "paragraphs" in partial_json:
                merged_paragraphs.extend(partial_json["paragraphs"])
            else:
                print("[WARN] No paragraphs or JSON parse error for chunk", i)

        final_doc_json = {"paragraphs": merged_paragraphs}

        try:
            self.save_docx_json(final_doc_json, out_path)
            messagebox.showinfo("Success", "Book corrected!")
        except Exception as e:
            messagebox.showerror(
                "Error", f"Error saving .docx file:\n{str(e)}")

        print("[DEBUG] Completed correction.")

    ##############################
    # Read docx, chunking
    ##############################
    def read_docx(self, file_path):
        if docx is None:
            raise ImportError("python-docx not installed.")
        d = Document(file_path)
        paragraphs = [p.text.strip() for p in d.paragraphs if p.text.strip()]
        return paragraphs

    def create_chunks(self, paragraphs, max_tokens):
        chunks = []
        current_chunk = []
        current_tokens = 0

        for para in paragraphs:
            if not para:
                continue
            p_tokens = self.count_tokens(para)
            if current_tokens + p_tokens > max_tokens and current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = [para]
                current_tokens = p_tokens
            else:
                current_chunk.append(para)
                current_tokens += p_tokens

        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        return chunks

    def count_tokens(self, text):
        try:
            enc = tiktoken.encoding_for_model("gpt-4")
        except KeyError:
            enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text))

    ##############################
    # Correction JSON
    ##############################
    def build_developer_prompt(self):
        tasks = []
        if self.fix_grammar.get():
            tasks.append("grammar")
        if self.fix_typos.get():
            tasks.append("typos")
        if self.fix_formatting.get():
            tasks.append("formatting")

        prompt = (
            f"You are a helpful assistant focusing on book/novel style corrections in {
                self.selected_language.get()}. "
            "If the text is not a novel, adapt corrections appropriately. "
            f"Fix only these aspects: {', '.join(tasks)}. "
            "Do NOT add or remove meaning. Keep the author's style. "
            "Focus on punctuation, spelling, capitalization, paragraph structuring. "
            "If it's a novel, ensure dialogue and chapter headings are properly formatted. "
        )
        extra = self.extra_instructions.get().strip()
        if extra:
            prompt += f"\nAdditional instructions:\n{extra}\n"
        return prompt

    def correct_chunk_json_mode(self, chunk, api_key):
        dev_prompt = self.build_developer_prompt()
        dev_instructions = (
            f"{dev_prompt}\n\n"
            "Do NOT add new content. Return valid JSON only. The JSON must have this structure:\n"
            "{\n"
            "   \"paragraphs\": [\n"
            "       { \"text\": \"corrected paragraph...\", \"style\": \"normal|heading|italic|bold|quote\" },\n"
            "       ...\n"
            "   ]\n"
            "}\n\n"
            "No extra keys. Only output valid JSON.\n"
            "Now correct the following text:\n\n"
        )

        messages = [
            {"role": "developer", "content": dev_instructions},
            {"role": "user", "content": chunk},
        ]

        data = {
            "model": "gpt-4o-mini",
            "messages": messages,
            "temperature": 0.0,
            "presence_penalty": 0.0,
            "frequency_penalty": 0.0,
            "response_format": {
                "type": "json_object"
            }
        }

        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        }

        try:
            response = requests.post(
                CHAT_COMPLETIONS_URL, headers=headers, json=data)
            response.raise_for_status()
            content = response.json()["choices"][0]["message"]["content"]
            parsed = json.loads(content.strip())
            return parsed
        except requests.exceptions.HTTPError as http_err:
            # Catch 401/403 etc., show a message
            print("[DEBUG] HTTPError:", http_err)
            messagebox.showerror("Error", f"API returned an error:\n{
                                 str(http_err)}\n\nCheck your API key.")
            return None
        except Exception as e:
            print("[DEBUG] Unexpected error in correct_chunk_json_mode:", e)
            return None

    ##############################
    # Save final docx
    ##############################
    def save_docx_json(self, doc_json, file_path):
        if docx is None:
            raise ImportError("python-docx not installed.")
        d = Document()

        style_normal = d.styles["Normal"]
        style_normal.font.name = "Times New Roman"
        style_normal.font.size = Pt(12)

        for p in doc_json.get("paragraphs", []):
            txt = p.get("text", "").strip()
            st = p.get("style", "normal")
            if not txt:
                continue
            paragraph = d.add_paragraph()
            run = paragraph.add_run(txt)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Force black text

            if st == "heading":
                paragraph.style = d.styles["Heading1"]
            elif st == "italic":
                run.italic = True
            elif st == "bold":
                run.bold = True
            elif st == "quote":
                if "Intense Quote" in d.styles:
                    paragraph.style = d.styles["Intense Quote"]
                else:
                    paragraph.style = d.styles["Quote"]
            else:
                paragraph.style = d.styles["Normal"]

        d.save(file_path)


def main():
    root = tk.Tk()
    app = AICorrectorGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
